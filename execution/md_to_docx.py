"""Convert CV markdown to styled DOCX with HCCH blue accents."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

HCCH_BLUE = RGBColor(0x00, 0x33, 0x66)
BODY_FONT = "Times New Roman"
INPUT_PATH = Path(__file__).parent.parent / "data" / "CV_Jisung_Kang(Mar2026).md"
OUTPUT_PATH = INPUT_PATH.with_suffix(".docx")


def add_blue_line(doc):
    """Add a thin HCCH blue horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="003366"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_formatted_runs(paragraph, text, default_size=Pt(11)):
    """Parse markdown inline formatting (**bold**, *italic*) into runs."""
    # Split by bold and italic patterns
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.name = BODY_FONT
        run.font.size = default_size


def add_bullet(doc, text, indent_level=0):
    """Add a bullet point with markdown formatting."""
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)

    # Handle links: → https://...
    link_match = re.match(r'(.*?)(→\s*https?://\S+)(.*)', text)
    if link_match:
        add_formatted_runs(p, link_match.group(1))
        run = p.add_run(link_match.group(2))
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
        run.font.color.rgb = HCCH_BLUE
        if link_match.group(3):
            add_formatted_runs(p, link_match.group(3))
    else:
        add_formatted_runs(p, text)


def build_docx():
    md = INPUT_PATH.read_text(encoding="utf-8")
    lines = md.split('\n')

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)

    # List Bullet style
    if 'List Bullet' in [s.name for s in doc.styles]:
        lb = doc.styles['List Bullet']
        lb.font.name = BODY_FONT
        lb.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # H1 - Name
        if line.startswith('# ') and not line.startswith('##'):
            name = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = BODY_FONT
            run.font.color.rgb = HCCH_BLUE
            i += 1
            continue

        # Contact line (right after name, before first ---)
        if i > 0 and not line.startswith('#') and not line.startswith('---') and line and not line.startswith('-'):
            # Check if we're in the header area (before first section)
            prev_lines = lines[:i]
            has_h2 = any(l.strip().startswith('## ') for l in prev_lines)
            if not has_h2:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                # Clean up markdown link syntax and pipes
                contact = line.replace('|', '  |  ').strip()
                contact = contact.rstrip('|').strip()
                run = p.add_run(contact)
                run.font.name = BODY_FONT
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                i += 1
                continue

        # Horizontal rule
        if line.startswith('---'):
            add_blue_line(doc)
            i += 1
            continue

        # H2 - Section heading
        if line.startswith('## '):
            title = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = BODY_FONT
            run.font.color.rgb = HCCH_BLUE
            # Add underline bar
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="003366"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # H3 - Sub-section heading
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.name = BODY_FONT
            run.font.color.rgb = HCCH_BLUE
            i += 1
            continue

        # Bullet items
        if line.startswith('- '):
            text = line[2:].strip()

            # Check if next lines are indented continuation (description lines)
            desc_lines = []
            j = i + 1
            while j < len(lines) and lines[j].startswith('  ') and not lines[j].strip().startswith('- '):
                desc_lines.append(lines[j].strip())
                j += 1

            add_bullet(doc, text)

            # Add description as indented follow-up
            for desc in desc_lines:
                if desc.startswith('→'):
                    # Link line
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(1)
                    run = p.add_run(desc)
                    run.font.name = BODY_FONT
                    run.font.size = Pt(10)
                    run.font.color.rgb = HCCH_BLUE
                elif desc:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(1)
                    add_formatted_runs(p, desc, Pt(10.5))

            i = j
            continue

        # Empty line - skip
        if not line.strip():
            i += 1
            continue

        # Default: plain paragraph
        p = doc.add_paragraph()
        add_formatted_runs(p, line)
        i += 1

    doc.save(str(OUTPUT_PATH))
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_docx()
